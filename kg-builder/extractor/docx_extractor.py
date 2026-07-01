# Extraction du DOCX de Lund (2018)
from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pycountry  # type: ignore[import-untyped]
from docx import Document  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]  # noqa: F401 — utilisé dans hyperlinks (activable)

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from config import settings  # noqa: E402


# Modèles de données

@dataclass
class DefinitionRecord:
    title_1:      str = ""
    title_2:      str = ""
    title_3:      str = ""
    title_4:      str = ""
    bold_terms:   str = ""
    sources:      str = ""
    urls:         list[str] = field(default_factory=list)
    references:   list[str] = field(default_factory=list)
    definition:   str = ""
    organization: str = ""
    country:      str = ""
    year:         str = ""

    def to_dict(self) -> dict:
        return {
            "title_1":      self.title_1,
            "title_2":      self.title_2,
            "title_3":      self.title_3,
            "title_4":      self.title_4,
            "bold_terms":   self.bold_terms,
            "sources":      self.sources,
            "urls":         "; ".join(self.urls),
            "references":   "; ".join(self.references),
            "definition":   self.definition,
            "organization": self.organization,
            "country":      self.country,
            "year":         self.year,
        }


@dataclass
class Table3Record:
    countries:            str = ""
    definition_type:      str = ""
    area_ha:              str = ""
    crown_cover_percent:  str = ""
    tree_height_m:        str = ""
    strip_width_m:        str = ""
    notes:                str = ""

    @property
    def is_unfccc(self) -> bool:
        return bool(re.search(r'\bUNFCCC\b', self.countries, re.IGNORECASE))

    def to_dict(self) -> dict:
        return {
            "countries":           self.countries,
            "definition_type":     self.definition_type,
            "area_ha":             self.area_ha,
            "crown_cover_percent": self.crown_cover_percent,
            "tree_height_m":       self.tree_height_m,
            "strip_width_m":       self.strip_width_m,
            "notes":               self.notes,
        }


# Extracteurs unitaires — logique spécifique à ce DOCX

def _bold_ext(p) -> str:
    """
    Extrait le terme initial en gras.

    Règle clé : le PREMIER run non vide doit être gras. Un paragraphe entier
    en gras (> 5 mots) n'est pas un terme-concept, il est ignoré.
    """
    runs = p.runs
    if not runs:
        return ""

    # Le premier run non vide doit être gras — sinon ce n'est pas un terme
    for run in runs:
        if run.text.strip():
            if not run.bold:
                return ""
            break
    else:
        return ""

    # Collecter les runs gras consécutifs (s'arrête dès qu'un run non-gras non vide est trouvé)
    bold_parts = []
    for run in runs:
        if not run.text.strip():
            continue
        if run.bold:
            bold_parts.append(run.text.strip())
        else:
            break

    bold_text = re.sub(r"\s+", " ", " ".join(bold_parts))
    bold_text = re.sub(r"^[\s\-–—:;.,/]+|[\s\-–—:;.,/]+$", "", bold_text)

    words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ\-]+", bold_text)
    return " ".join(words) if 1 <= len(words) <= 5 else ""


def _term_ext(p) -> str:
    """
    Extrait le terme dans le pattern :  (source) TERME - définition
                                     ou (source) TERME -  définition
                                     ou (source) TERME :  définition

    Séparateurs exacts de ce DOCX : '--' | '-' | ':'
    Max 5 mots.
    """
    m = re.match(r"^\([^)]*\)\s+(.+?)\s+(--|-|:)\s+", p.text or "")
    if not m:
        return ""
    words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ\-]+", m.group(1).strip())
    return " ".join(words) if 1 <= len(words) <= 5 else ""


def _parentheses_ext(text: str) -> list[str]:
    """Extrait toutes les parties entre parenthèses (sources, années, organisations)."""
    return re.findall(r'\(([^()]+)\)', text)


def _urls_ext(p, doc) -> list[str]:
    """
    Extrait les URLs visibles du paragraphe.

    Le paramètre doc est conservé pour l'activation future de l'extraction
    des URLs cachées dans les hyperliens Word (w:hyperlink via doc.part.rels).
    Pour activer, décommenter le bloc ci-dessous.
    """
    url_list: list[str] = []

    # Hyperliens cachés (à activer si nécessaire)
    # for hyperlink in p._p.xpath('.//w:hyperlink'):
    #     rId = hyperlink.get(qn('r:id'))
    #     if rId and rId in doc.part.rels:
    #         url = getattr(doc.part.rels[rId], "target", None)
    #         if url:
    #             url_list.append(url)

    # URLs visibles dans le texte
    for u in re.findall(r"https?://[A-Za-z0-9\-._~:/?#\[\]@!$&'()*+,;=%]+", p.text):
        url_list.append(u)

    # Nettoyage + déduplication (ordre préservé)
    seen:   set[str]  = set()
    result: list[str] = []
    for u in url_list:
        u = u.rstrip(".,;:!)]}")
        if u not in seen:
            seen.add(u)
            result.append(u)
    return result


def _references_ext(p) -> list[str]:
    """Extrait les références 'Ref. ...' du paragraphe."""
    refs = []
    for m in re.finditer(r'(?i)(?<!\w)Ref\.?\s+(.+)', (p.text or "").strip()):
        ref = "Ref. " + re.sub(r'[\s;:,\./]+$', '', m.group(1).strip())
        refs.append(ref)
    return list(dict.fromkeys(refs))


def _definition_ext(text: str, bold_term: str, term: str,
                    refs: list[str], urls: list[str], source: str) -> str:
    """
    Nettoie le texte de définition.
    Supprime dans l'ordre : URLs → refs → source entre () → bold_term → term.
    """
    clean = text
    for u in urls:
        clean = clean.replace(u, "")
    for r in refs:
        clean = clean.replace(r, "")
    if source:
        clean = clean.replace(f"({source})", "", 1)
    if bold_term:
        clean = clean.replace(bold_term, "", 1)
    if term:
        clean = clean.replace(term, "", 1)

    clean = re.sub(r"\s+", " ", clean)
    clean = re.sub(r"^[\s\n\r\t,;:\-=–—./()]+|[\s\n\r\t,;:\-=–—./()]+$", "", clean)
    return clean.strip()


def _year_ext(s: str) -> str:
    m = re.search(r"(19|20)\d{2}", s)
    return m.group(0) if m else ""


# La note se découpe différemment selon la section (scope via title_3) :
# International/General = juste une organisation, jamais de pays.
# National = pays + organisation + année, sauf "USA-FED-..." qui reste USA.
# State/local = "Pays-Région", sauf "USA-STATE-Georgia" (sinon confondu
# avec le pays Géorgie).
def _find_country_match(s: str) -> tuple[str, int, int] | None:
    """
    Trouve un seul nom de pays dans s, avec sa position dans le texte.
    Une note peut citer 2 pays ("Serbia & Montenegro") : il faut en garder
    un seul, sinon _org_ext retire les deux et il reste juste "&".
    """
    t = s.lower()
    for country in pycountry.countries:
        name = country.name.lower()
        idx = t.find(name)
        if idx != -1:
            return country.name, idx, idx + len(name)
    return None


def _country_ext(s: str, scope: str = "") -> str:
    t = s.strip().lower()
    if t.startswith("usa-state") or t.startswith("usa-fed"):
        return "United States"
    if scope in ("International", "General"):
        return ""
    match = _find_country_match(s)
    return match[0] if match else ""


def _org_ext(s: str, scope: str = "") -> str:
    """Extrait l'organisation = résidu après suppression de l'année et du pays."""
    t = s.strip().lower()
    if t.startswith("usa-state") or t.startswith("usa-fed"):
        org = re.sub(r"(19|20)\d{2}", "", s)
        return org.strip(" -;.&")
    org = s
    if scope not in ("International", "General"):
        match = _find_country_match(s)
        if match:
            _, start, end = match
            org = org[:start] + org[end:]
    org = re.sub(r"(19|20)\d{2}", "", org).strip(" -;.&")
    # parfois la note c'est juste une mesure ("10 ft", "45°"), pas une organisation
    return org if re.search(r"[A-Za-z]", org) else ""


# Arbre de titres

def _get_heading_level(p) -> Optional[int]:
    style = (p.style.name or "").lower()
    if style.startswith("heading"):
        try:
            return int(style.split()[-1])
        except (ValueError, IndexError):
            return 1
    return None


def _clean_title(text: str) -> str:
    return re.sub(r'^\d+(\.\d+)*\s*', '', text).strip()


def _build_tree(doc: Document) -> dict:
    root  = {"title": None, "level": 0, "children": [], "paragraphs": []}
    stack = [root]
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        level = _get_heading_level(p)
        if level is not None:
            node = {"title": _clean_title(text), "level": level,
                    "children": [], "paragraphs": []}
            while stack and stack[-1]["level"] >= level:
                stack.pop()
            stack[-1]["children"].append(node)
            stack.append(node)
        else:
            stack[-1]["paragraphs"].append(p)
    return root


# Traversée et extraction

_JUNK_RE = re.compile(r'^([-–—]+|\d+(\.\d+)*|\(\s*\))$')


def _traverse(node: dict, doc: Document,
              parents: list[str], records: list[DefinitionRecord]) -> None:
    current = (parents + [node["title"]]) if node["title"] else parents

    for p in node["paragraphs"]:
        if not p.text.strip():
            continue

        bold_term = _bold_ext(p)
        term      = _term_ext(p)

        sources   = _parentheses_ext(p.text)
        source    = sources[0] if sources else ""

        urls  = _urls_ext(p, doc)
        refs  = _references_ext(p)
        defn  = _definition_ext(p.text, bold_term, term, refs, urls, source)

        # Sauter définitions vides ou junk (tiret seul, chiffre seul, () vide)
        if not defn or _JUNK_RE.fullmatch(defn):
            continue

        padded = (current + ["", "", "", ""])[:4]
        scope  = settings.scope_map.get(padded[2], "")
        records.append(DefinitionRecord(
            title_1=padded[0],      title_2=padded[1],
            title_3=padded[2],      title_4=padded[3],
            bold_terms="; ".join(x for x in [bold_term, term] if x),
            sources=source,
            urls=urls,
            references=refs,
            definition=defn,
            organization=_org_ext(source, scope),
            country=_country_ext(source, scope),
            year=_year_ext(source),
        ))

    for child in node["children"]:
        _traverse(child, doc, current, records)


# Interface publique

def extract_definitions(docx_path: str | Path) -> list[DefinitionRecord]:
    doc  = Document(str(docx_path))
    tree = _build_tree(doc)
    records: list[DefinitionRecord] = []
    _traverse(tree, doc, [], records)
    return records


def extract_table3(docx_path: str | Path) -> list[Table3Record]:
    """
    Extrait la Table 3 'National criteria used for defining forestland'.
    Cherche le tableau par son contenu, lit les 7 colonnes fixes.
    """
    doc = Document(str(docx_path))
    target = None
    for table in doc.tables:
        text = " ".join(c.text.lower() for row in table.rows for c in row.cells)
        if "national criteria used for defining forestland" in text:
            target = table
            break
    if not target:
        return []

    _HEADERS = {"country", "countries", "definition type", "area", "crown", "height", "strip", "notes"}
    records: list[Table3Record] = []

    for row in target.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) != 7:
            continue
        cleaned = [re.sub(r"\s+", " ", c).strip(" -\n\t") for c in cells]
        if not cleaned[0] or any(h in cleaned[0].lower() for h in _HEADERS):
            continue
        records.append(Table3Record(
            countries=cleaned[0],   definition_type=cleaned[1],
            area_ha=cleaned[2],     crown_cover_percent=cleaned[3],
            tree_height_m=cleaned[4], strip_width_m=cleaned[5],
            notes=cleaned[6],
        ))
    return records
